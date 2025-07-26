
# from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
# from sqlalchemy.orm import Session
# from typing import List
# import magic  # python-magic for file type detection
# import chardet  # for character encoding detection
# import PyPDF2  # for PDF parsing
# import docx  # for Word document parsing
# import io
# from app.core.database import get_db
# from app.core.auth import get_current_user
# from app.models.user import User
# from app.models.document import Document
# from app.schemas.document import DocumentCreate, DocumentResponse, DocumentList
# from app.services.rag_service import RAGService
# from loguru import logger

# router = APIRouter(prefix="/documents", tags=["Documents"])
# rag_service = RAGService()

# # File size limit (10MB)
# MAX_FILE_SIZE = 10 * 1024 * 1024
# ALLOWED_CONTENT_TYPES = [
#     "text/plain",
#     "text/csv",
#     "application/json",
#     "text/markdown",
#     "application/pdf",
#     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
#     "application/msword"  # .doc
# ]

# @router.post("/", response_model=DocumentResponse)
# async def create_document(
#     document_data: DocumentCreate,
#     current_user: User = Depends(get_current_user),
#     db: Session = Depends(get_db)
# ):
#     """Create a new document and add it to the RAG system"""
#     try:
#         # Validate input
#         if not document_data.title.strip():
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="Document title cannot be empty"
#             )
        
#         if not document_data.content.strip():
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="Document content cannot be empty"
#             )
        
#         # Create document in database
#         db_document = Document(
#             title=document_data.title.strip(),
#             content=document_data.content,
#             owner_id=current_user.id,
#             pinecone_id=""  # Will be updated by RAG service
#         )
        
#         db.add(db_document)
#         db.flush()  # Get the ID without committing
        
#         # Add to RAG system
#         await rag_service.add_document(db_document, db)
        
#         # Commit the transaction
#         db.commit()
        
#         logger.info(f"Created document: {document_data.title}")
#         return DocumentResponse.from_orm(db_document)
        
#     except HTTPException:
#         db.rollback()
#         raise
#     except Exception as e:
#         logger.error(f"Failed to create document: {e}")
#         db.rollback()
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Failed to create document: {str(e)}"
#         )

# @router.get("/", response_model=List[DocumentList])
# async def list_documents(
#     current_user: User = Depends(get_current_user),
#     db: Session = Depends(get_db)
# ):
#     """List all documents for the current user"""
#     documents = db.query(Document).filter(Document.owner_id == current_user.id).all()
#     return [DocumentList.from_orm(doc) for doc in documents]

# @router.get("/{document_id}", response_model=DocumentResponse)
# async def get_document(
#     document_id: int,
#     current_user: User = Depends(get_current_user),
#     db: Session = Depends(get_db)
# ):
#     """Get a specific document"""
#     document = db.query(Document).filter(
#         Document.id == document_id,
#         Document.owner_id == current_user.id
#     ).first()
    
#     if not document:
#         raise HTTPException(
#             status_code=status.HTTP_404_NOT_FOUND,
#             detail="Document not found"
#         )
    
#     return DocumentResponse.from_orm(document)

# @router.delete("/{document_id}")
# async def delete_document(
#     document_id: int,
#     current_user: User = Depends(get_current_user),
#     db: Session = Depends(get_db)
# ):
#     """Delete a document from the system"""
#     document = db.query(Document).filter(
#         Document.id == document_id,
#         Document.owner_id == current_user.id
#     ).first()
    
#     if not document:
#         raise HTTPException(
#             status_code=status.HTTP_404_NOT_FOUND,
#             detail="Document not found"
#         )
    
#     try:
#         await rag_service.delete_document(document, db)
#         db.commit()
#         return {"message": "Document deleted successfully"}
#     except Exception as e:
#         logger.error(f"Failed to delete document: {e}")
#         db.rollback()
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Failed to delete document: {str(e)}"
#         )

# def extract_text_from_pdf(content: bytes) -> str:
#     """Extract text from PDF content"""
#     try:
#         pdf_file = io.BytesIO(content)
#         pdf_reader = PyPDF2.PdfReader(pdf_file)
#         text = ""
        
#         for page in pdf_reader.pages:
#             text += page.extract_text() + "\n"
        
#         return text.strip()
#     except Exception as e:
#         logger.error(f"Failed to extract text from PDF: {e}")
#         raise HTTPException(
#             status_code=status.HTTP_400_BAD_REQUEST,
#             detail="Failed to extract text from PDF file"
#         )

# def extract_text_from_docx(content: bytes) -> str:
#     """Extract text from DOCX content"""
#     try:
#         doc_file = io.BytesIO(content)
#         doc = docx.Document(doc_file)
#         text = ""
        
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"
        
#         return text.strip()
#     except Exception as e:
#         logger.error(f"Failed to extract text from DOCX: {e}")
#         raise HTTPException(
#             status_code=status.HTTP_400_BAD_REQUEST,
#             detail="Failed to extract text from Word document"
#         )

# def extract_text_from_file(content: bytes, file_type: str) -> str:
#     """Extract text from various file types"""
#     if file_type == "application/pdf":
#         return extract_text_from_pdf(content)
#     elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
#         return extract_text_from_docx(content)
#     else:
#         # Handle text-based files
#         try:
#             # First try UTF-8
#             return content.decode('utf-8')
#         except UnicodeDecodeError:
#             # If UTF-8 fails, detect encoding
#             detected = chardet.detect(content)
#             encoding = detected.get('encoding', 'utf-8')
#             confidence = detected.get('confidence', 0)
            
#             if confidence < 0.7:
#                 raise HTTPException(
#                     status_code=status.HTTP_400_BAD_REQUEST,
#                     detail="Unable to detect file encoding with sufficient confidence"
#                 )
            
#             try:
#                 return content.decode(encoding)
#             except UnicodeDecodeError:
#                 raise HTTPException(
#                     status_code=status.HTTP_400_BAD_REQUEST,
#                     detail="Unable to decode file content"
#                 )

# @router.post("/upload")
# async def upload_document(
#     title: str = Form(...),
#     file: UploadFile = File(...),
#     current_user: User = Depends(get_current_user),
#     db: Session = Depends(get_db)
# ):
#     """Upload a document file"""
#     try:
#         # Validate title
#         if not title.strip():
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="Document title cannot be empty"
#             )
        
#         # Validate file
#         if not file.filename:
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="No file provided"
#             )
        
#         # Check file size
#         content = await file.read()
#         if len(content) > MAX_FILE_SIZE:
#             raise HTTPException(
#                 status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
#                 detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
#             )
        
#         # Detect file type
#         file_type = magic.from_buffer(content, mime=True)
#         if file_type not in ALLOWED_CONTENT_TYPES:
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail=f"Unsupported file type: {file_type}. Allowed types: {', '.join(['PDF', 'DOCX', 'TXT', 'CSV', 'JSON', 'Markdown'])}"
#             )
        
#         # Extract text content based on file type
#         text_content = extract_text_from_file(content, file_type)
        
#         # Validate content
#         if not text_content.strip():
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="File content is empty or could not be extracted"
#             )
        
#         # Create document in database
#         db_document = Document(
#             title=title.strip(),
#             content=text_content,
#             file_path=file.filename,
#             owner_id=current_user.id,
#             pinecone_id=""
#         )
        
#         db.add(db_document)
#         db.flush()
        
#         # Add to RAG system
#         await rag_service.add_document(db_document, db)
        
#         # Commit the transaction
#         db.commit()
        
#         logger.info(f"Uploaded document: {title} (file: {file.filename}, type: {file_type})")
#         return {
#             "message": "Document uploaded successfully", 
#             "document_id": db_document.id,
#             "title": db_document.title,
#             "filename": file.filename,
#             "file_type": file_type,
#             "content_length": len(text_content)
#         }
        
#     except HTTPException:
#         db.rollback()
#         raise
#     except Exception as e:
#         logger.error(f"Failed to upload document: {e}")
#         db.rollback()
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Failed to upload document: {str(e)}"
#         )

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List
import magic  # python-magic for file type detection
import chardet  # for character encoding detection
import PyPDF2  # for PDF parsing
import docx  # for Word document parsing
import io
from app.core.database import get_db
from app.core.auth import get_current_user
from app.models.user import User
from app.models.document import Document
from app.schemas.document import DocumentCreate, DocumentResponse, DocumentList
from app.services.rag_service import RAGService
from loguru import logger

router = APIRouter(prefix="/documents", tags=["Documents"])
rag_service = RAGService()

# File size limit (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024
ALLOWED_CONTENT_TYPES = [
    "text/plain",
    "text/csv",
    "application/json",
    "text/markdown",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword"  # .doc
]

@router.post("/", response_model=DocumentResponse)
async def create_document(
    document_data: DocumentCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new document and add it to the RAG system"""
    try:
        # Validate input
        if not document_data.title.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document title cannot be empty"
            )
        
        if not document_data.content.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document content cannot be empty"
            )
        
        # Create document in database
        db_document = Document(
            title=document_data.title.strip(),
            content=document_data.content,
            owner_id=current_user.id,
            pinecone_id=""  # Will be updated by RAG service
        )
        
        db.add(db_document)
        db.flush()  # Get the ID without committing
        
        # Add to RAG system
        rag_service.add_document(db_document, db)
        
        # Commit the transaction
        db.commit()
        
        logger.info(f"Created document: {document_data.title}")
        return DocumentResponse.from_orm(db_document)
        
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        logger.error(f"Failed to create document: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to create document: {str(e)}"
        )

@router.get("/", response_model=List[DocumentList])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all documents for the current user"""
    documents = db.query(Document).filter(Document.owner_id == current_user.id).all()
    return [DocumentList.from_orm(doc) for doc in documents]

@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific document"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.owner_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return DocumentResponse.from_orm(document)

@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a document from the system"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.owner_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        rag_service.delete_document(document, db)
        db.commit()
        return {"message": "Document deleted successfully"}
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content"""
    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from PDF file"
        )

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content"""
    try:
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from Word document"
        )

def extract_text_from_file(content: bytes, file_type: str) -> str:
    """Extract text from various file types"""
    if file_type == "application/pdf":
        return extract_text_from_pdf(content)
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        return extract_text_from_docx(content)
    else:
        # Handle text-based files
        try:
            # First try UTF-8
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # If UTF-8 fails, detect encoding
            detected = chardet.detect(content)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)
            
            if confidence < 0.7:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Unable to detect file encoding with sufficient confidence"
                )
            
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Unable to decode file content"
                )

@router.post("/upload")
async def upload_document(
    title: str = Form(...),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a document file"""
    try:
        # Validate title
        if not title.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document title cannot be empty"
            )
        
        # Validate file
        if not file.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No file provided"
            )
        
        # Check file size
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # Detect file type
        file_type = magic.from_buffer(content, mime=True)
        if file_type not in ALLOWED_CONTENT_TYPES:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_type}. Allowed types: {', '.join(['PDF', 'DOCX', 'TXT', 'CSV', 'JSON', 'Markdown'])}"
            )
        
        # Extract text content based on file type
        text_content = extract_text_from_file(content, file_type)
        
        # Validate content
        if not text_content.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File content is empty or could not be extracted"
            )
        
        # Create document in database
        db_document = Document(
            title=title.strip(),
            content=text_content,
            file_path=file.filename,
            owner_id=current_user.id,
            pinecone_id=""
        )
        
        db.add(db_document)
        db.flush()
        
        # Add to RAG system
        rag_service.add_document(db_document, db)
        
        # Commit the transaction
        db.commit()
        
        logger.info(f"Uploaded document: {title} (file: {file.filename}, type: {file_type})")
        return {
            "message": "Document uploaded successfully", 
            "document_id": db_document.id,
            "title": db_document.title,
            "filename": file.filename,
            "file_type": file_type,
            "content_length": len(text_content)
        }
        
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        logger.error(f"Failed to upload document: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload document: {str(e)}"
        )